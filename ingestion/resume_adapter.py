import os
import re
from typing import List, Dict, Any
from ingestion.base import SourceAdapter, RawRecord

# Try loading pdf/docx, if fail, we handle gracefully
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

class ResumeAdapter(SourceAdapter):
    def detect(self, filepath_or_url: str) -> bool:
        if not os.path.exists(filepath_or_url):
            return False
        ext = filepath_or_url.lower().split('.')[-1]
        return ext in ['pdf', 'docx']

    def _extract_pdf_text(self, filepath: str) -> str:
        # Try pdfplumber first as it preserves layout and paragraph line breaks
        if pdfplumber is not None:
            try:
                with pdfplumber.open(filepath) as pdf:
                    text = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    if text:
                        return "\n".join(text)
            except Exception:
                pass

        # Fallback to pypdf
        if PdfReader is not None:
            try:
                reader = PdfReader(filepath)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
            except Exception:
                pass

        return ""

    def _extract_docx_text(self, filepath: str) -> str:
        if docx is None:
            return ""
        try:
            doc = docx.Document(filepath)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return "\n".join(text)
        except Exception:
            return ""

    def extract(self, filepath_or_url: str) -> List[RawRecord]:
        records = []
        if not os.path.exists(filepath_or_url):
            return records

        ext = filepath_or_url.lower().split('.')[-1]
        text = ""
        
        # 1. Try LlamaParse for PDF first if available
        if ext == 'pdf':
            try:
                from ingestion.ai_parser import parse_pdf_with_llama_parse
                text = parse_pdf_with_llama_parse(filepath_or_url)
            except Exception:
                pass
                
        # 2. Fallback to local text extraction (pdfplumber/docx)
        if not text.strip():
            if ext == 'pdf':
                text = self._extract_pdf_text(filepath_or_url)
            elif ext == 'docx':
                text = self._extract_docx_text(filepath_or_url)

        if not text.strip():
            return records

        source_name = "Resume"
        record_id = os.path.basename(filepath_or_url)

        # 3. Try AI-based structured extraction using Mistral
        try:
            from ingestion.ai_parser import extract_structured_data_with_mistral
            ai_records = extract_structured_data_with_mistral(text, source_name, record_id)
            if ai_records:
                return ai_records
        except Exception:
            pass

        # 4. Fallback to local heuristic-based extraction
        method_name = "inference"
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        name = None
        for line in lines:
            # Skip lines that are too long, contain symbols, or look like contact info
            if len(line) < 5 or len(line) > 50:
                continue
            if "@" in line or "/" in line or "resume" in line.lower() or "page" in line.lower():
                continue
            # Basic alphabetic check (allowing spaces and dots)
            if re.match(r'^[A-Za-z\s\.\-\’]+$', line):
                name = line
                break
        
        if name:
            records.append(RawRecord(field="full_name", value=name, source=source_name, method=method_name, record_id=record_id))

        # 2. Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        for email in list(set(emails)):
            records.append(RawRecord(field="emails", value=email, source=source_name, method=method_name, record_id=record_id))

        # 3. Extract phones
        phone_pattern = r'\+?\b\d[\d\-\(\)\s\.]{7,16}\d\b'
        phones = re.findall(phone_pattern, text)
        # Filter out numbers that look like dates or page numbers
        valid_phones = []
        for p in phones:
            digits = re.sub(r'\D', '', p)
            if len(digits) >= 9 and len(digits) <= 15:
                valid_phones.append(p.strip())
        for phone in list(set(valid_phones)):
            records.append(RawRecord(field="phones", value=phone, source=source_name, method=method_name, record_id=record_id))

        # 4. Extract links (LinkedIn, GitHub)
        linkedin_pattern = r'https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9_\-\u00C0-\u00FF]+'
        github_pattern = r'https?://(?:www\.)?github\.com/[A-Za-z0-9_\-]+'
        
        linkedins = re.findall(linkedin_pattern, text, re.IGNORECASE)
        githubs = re.findall(github_pattern, text, re.IGNORECASE)

        links_dict = {}
        if linkedins:
            links_dict["linkedin"] = linkedins[0]
        if githubs:
            links_dict["github"] = githubs[0]
            
        if links_dict:
            records.append(RawRecord(field="links", value=links_dict, source=source_name, method=method_name, record_id=record_id))

        # 5. Extract Skills, Experience, and Education using Section Parsing
        sections = self._parse_sections(text)
        
        # Parse skills
        skills = self._parse_skills(sections.get("skills", ""))
        for skill in skills:
            records.append(RawRecord(field="skills", value=skill, source=source_name, method=method_name, record_id=record_id))

        # Parse experience
        exp_entries = self._parse_experience(sections.get("experience", ""))
        for exp in exp_entries:
            records.append(RawRecord(field="experience", value=exp, source=source_name, method=method_name, record_id=record_id))

        # Parse education
        edu_entries = self._parse_education(sections.get("education", ""))
        for edu in edu_entries:
            records.append(RawRecord(field="education", value=edu, source=source_name, method=method_name, record_id=record_id))

        return records

    def _parse_sections(self, text: str) -> Dict[str, str]:
        # Divide the text into sections
        lines = text.split('\n')
        sections = {"skills": [], "experience": [], "education": [], "other": []}
        current_section = "other"

        # Keywords for sections
        section_keywords = {
            "skills": ["skills", "technical skills", "technologies", "expertise", "core competencies", "languages and tools"],
            "experience": ["experience", "work experience", "employment history", "professional experience", "work history"],
            "education": ["education", "academic history", "academic background", "credentials"]
        }

        for line in lines:
            line_clean = line.strip().lower()
            # Check if this line is a section header (short line and starts/ends with keywords or exact matches)
            is_header = False
            if len(line_clean) < 30:
                for sec, keywords in section_keywords.items():
                    for kw in keywords:
                        # Match whole word or exact header
                        if line_clean == kw or (kw in line_clean and len(line_clean) < len(kw) + 4):
                            current_section = sec
                            is_header = True
                            break
                    if is_header:
                        break
            if not is_header:
                sections[current_section].append(line)

        return {k: "\n".join(v) for k, v in sections.items()}

    def _parse_skills(self, section_text: str) -> List[str]:
        if not section_text.strip():
            return []
        
        # Split by typical delimiters: commas, bullets, pipes, or newlines
        delimiters = r'[,|•\t]|\n'
        tokens = re.split(delimiters, section_text)
        skills = []
        for t in tokens:
            cleaned = t.strip()
            # Basic validation of skill strings
            if cleaned and len(cleaned) > 1 and len(cleaned) < 30:
                # Exclude bullets and weird characters
                cleaned = re.sub(r'^[-\*•]\s*', '', cleaned).strip()
                if cleaned and not cleaned.lower().startswith("proficient") and not cleaned.lower().startswith("knowledge"):
                    skills.append(cleaned)
        return list(set(skills))

    def _parse_experience(self, section_text: str) -> List[Dict[str, Any]]:
        entries = []
        if not section_text.strip():
            return entries

        # We will split paragraphs. Usually a job has a company name, title, date range, and bullets.
        # Let's group lines. A new job entry often begins with a line containing a year/month-year range.
        date_pattern = r'\b(19|20)\d{2}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b|present'
        lines = [l.strip() for l in section_text.split('\n') if l.strip()]
        
        current_job = None
        
        for line in lines:
            # Check if line contains a date range (e.g. 2021 - Present or June 2020 - 2022)
            has_date = len(re.findall(date_pattern, line, re.IGNORECASE)) >= 2 or (("present" in line.lower() or "current" in line.lower()) and len(re.findall(date_pattern, line, re.IGNORECASE)) >= 1)
            
            if has_date and len(line) < 100:
                if current_job:
                    entries.append(current_job)
                
                # Try to parse company, title, date
                # Simple heuristic: split by comma, dash, pipe, or tab
                parts = re.split(r'[,|–\-\t]', line)
                title = parts[0].strip() if len(parts) > 0 else None
                company = parts[1].strip() if len(parts) > 1 else None
                
                # Find date range in the line
                date_matches = list(re.finditer(r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}|\b\d{4}|\bpresent\b', line, re.IGNORECASE))
                start_date = date_matches[0].group(0) if len(date_matches) > 0 else None
                end_date = date_matches[1].group(0) if len(date_matches) > 1 else ("Present" if "present" in line.lower() else None)
                
                current_job = {
                    "company": company or "Unknown Company",
                    "title": title or "Unknown Title",
                    "start": start_date,
                    "end": end_date,
                    "summary": ""
                }
            else:
                if current_job:
                    # Append description
                    sep = "\n" if current_job["summary"] else ""
                    current_job["summary"] += sep + line
                elif len(line) > 15 and len(line) < 100:
                    # Create an ad-hoc job if we haven't found a header yet
                    current_job = {
                        "company": "Company Reference",
                        "title": line,
                        "start": None,
                        "end": None,
                        "summary": ""
                    }
                    
        if current_job:
            entries.append(current_job)
            
        # Clean summaries
        for entry in entries:
            entry["summary"] = entry["summary"].strip()
            # Clean date strings
            if entry["start"]: entry["start"] = entry["start"].strip()
            if entry["end"]: entry["end"] = entry["end"].strip()
            
        return entries

    def _parse_education(self, section_text: str) -> List[Dict[str, Any]]:
        entries = []
        if not section_text.strip():
            return entries

        lines = [l.strip() for l in section_text.split('\n') if l.strip()]
        edu_keywords = ["university", "college", "school", "institute", "polytechnic"]
        degree_keywords = ["bachelor", "master", "phd", "b.s", "m.s", "b.a", "m.a", "diploma", "degree", "associate"]

        for line in lines:
            is_edu_line = any(kw in line.lower() for kw in edu_keywords + degree_keywords)
            if is_edu_line:
                # Find year
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                end_year = int(year_match.group(0)) if year_match else None
                
                # Split line to estimate school/degree
                parts = re.split(r'[,|\-\t]', line)
                inst = None
                deg = None
                major = None
                
                for part in parts:
                    part_clean = part.strip()
                    if any(kw in part_clean.lower() for kw in edu_keywords):
                        inst = part_clean
                    elif any(kw in part_clean.lower() for kw in degree_keywords):
                        deg = part_clean
                    else:
                        if not major and len(part_clean) > 3:
                            major = part_clean
                
                if not inst and len(parts) > 0:
                    inst = parts[0].strip()
                if not deg and len(parts) > 1:
                    deg = parts[1].strip()
                    
                entries.append({
                    "institution": inst or "Unknown Institution",
                    "degree": deg or "Degree",
                    "field": major,
                    "end_year": end_year
                })
        return entries
